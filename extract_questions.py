import openai
import requests
import fitz  # PyMuPDF for PDF handling
import shutil
import os
from datetime import datetime
from docx import Document as DocxDocument  # for handling DOCX files
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain.schema import Document
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
import logging

# Set OpenAI API key directly
openai.api_key = os.getenv('OPENAI_API_KEY')  # Replace with your actual API key

CHROMA_PATH = "chroma"
DATA_PATH = "data"
BATCH_SIZE = 1  

def identify_eligibility_section(page_text):
    """Identify if a given page belongs to the eligibility criteria section."""
    prompt = f"""
    The following is a page from an RFP document. Identify if this page belongs to the eligibility criteria section. Respond with 'yes' or 'no'.

    Page Text:
    {page_text}

    Does this page belong to the eligibility criteria section?
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=5,
        n=1,
        stop=None,
        temperature=0.7,
    )
    answer = response.choices[0].message['content'].strip().lower()
    return 'yes' in answer

def extract_text_from_eligibility_section_from_link(s3_link):
    """Extract text from the eligibility section of a document using an S3 link."""
    temp_file_path = "temp_file"
    try:
        response = requests.get(s3_link)
        response.raise_for_status()

        content_type = response.headers.get('Content-Type')

        if 'pdf' in content_type:
            temp_file_path += ".pdf"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(response.content)
            text_content = extract_text_from_pdf(temp_file_path)  # Extract text before deleting the file
        elif 'text' in content_type:
            temp_file_path += ".txt"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(response.content)
            with open(temp_file_path, "r") as temp_file:
                text_content = temp_file.read()
        elif 'word' in content_type or 'officedocument' in content_type:
            temp_file_path += ".docx"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(response.content)
            text_content = extract_text_from_docx(temp_file_path)
        else:
            raise Exception(f"Unsupported file type: {content_type}")

        return text_content

    except Exception as e:
        raise Exception(f"Error processing document from S3: {e}")

    finally:
        # Ensure the file is closed before trying to delete it
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except Exception as e:
                logging.error(f"Error deleting temporary file {temp_file_path}: {str(e)}")

def extract_full_text_from_link(s3_link):
    """Extract full text from a document using an S3 link."""
    temp_file_path = "temp_file"
    try:
        response = requests.get(s3_link)
        response.raise_for_status()

        content_type = response.headers.get('Content-Type')

        if 'pdf' in content_type:
            temp_file_path += ".pdf"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(response.content)
            text_content = extract_full_text_from_pdf(temp_file_path)
        elif 'text' in content_type:
            temp_file_path += ".txt"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(response.content)
            with open(temp_file_path, "r") as temp_file:
                text_content = temp_file.read()
        elif 'word' in content_type or 'officedocument' in content_type:
            temp_file_path += ".docx"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(response.content)
            text_content = extract_full_text_from_docx(temp_file_path)
        else:
            raise Exception(f"Unsupported file type: {content_type}")

        return text_content

    except Exception as e:
        raise Exception(f"Error processing document from S3: {e}")

    finally:
        # Ensure the file is closed before trying to delete it
        if os.path.exists(temp_file_path):
            try:
                os.remove(temp_file_path)
            except Exception as e:
                logging.error(f"Error deleting temporary file {temp_file_path}: {str(e)}")


def extract_text_from_pdf(pdf_path):
    """Extract text from PDF."""
    document = fitz.open(pdf_path)
    eligibility_text = ""

    eligibility_section_started = False
    for page_num in range(len(document)):
        page = document.load_page(page_num)
        page_text = page.get_text("text")

        if identify_eligibility_section(page_text):
            eligibility_section_started = True

        if eligibility_section_started:
            eligibility_text += page_text
            if not identify_eligibility_section(page_text) and len(eligibility_text.split()) > 500:
                break

    return eligibility_text.strip()

def extract_full_text_from_pdf(pdf_path):
    """Extract full text from PDF."""
    document = fitz.open(pdf_path)
    full_text = ""

    for page_num in range(len(document)):
        page = document.load_page(page_num)
        page_text = page.get_text("text")
        full_text += page_text

    return full_text.strip()

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file."""
    doc = DocxDocument(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_full_text_from_docx(docx_path):
    """Extract full text from DOCX file."""
    doc = DocxDocument(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_questions_from_eligibility_section(content):
    """Convert eligibility section content into a series of questions."""
    prompt = f"""
    Analyze the following eligibility criteria content from an RFP document and convert it into a series of questions that would help a potential bidder understand the requirements clearly. Make sure to cover all important points and details in the form of questions.

    Eligibility Criteria Content:
    {content}

    Convert the above content into questions.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
        n=1,
        stop=None,
        temperature=0.7,
    )
    return response['choices'][0].message['content'].strip()

def generate_questions_from_text(text):
    """Generate questions from any given text."""
    text = text.strip()  # Ensure the text is clean and stripped
    
    prompt = f"""
    The following is a section from an RFP document. Please extract and convert the content into a series of questions that would help a potential bidder understand the requirements clearly. Ensure to cover all the key points and details in the form of questions.

    RFP Content:
    {text}

    Convert the above content into questions.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=500,
        n=1,
        stop=None,
        temperature=0.7,
    )

    cleaned_questions = response['choices'][0].message['content'].strip()

    # Ensure to split the generated questions properly if needed
    return [q.strip() for q in cleaned_questions.split("\n") if q.strip()]

def generate_content_using_prompt(content, prompt):
    """Generate content using a provided prompt."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": f"{prompt}\n\n{content}"}
        ],
        max_tokens=1024,
        n=1,
        stop=None,
        temperature=0.7,
    )
    return response['choices'][0].message['content'].strip()

def main():
    print("[DEBUG] Starting script...")
    print(f"[DEBUG] OpenAI API Key: {openai.api_key}")

    print("[DEBUG] Loading documents...")
    documents = load_documents()
    print(f"[DEBUG] Loaded {len(documents)} documents")

    for batch_start in range(0, len(documents), BATCH_SIZE):
        batch_end = min(batch_start + BATCH_SIZE, len(documents))
        batch = documents[batch_start:batch_end]
        print(f"[DEBUG] Processing batch {batch_start + 1} to {batch_end}...")

        for document in batch:
            try:
                print(f"[DEBUG] Processing document {document.metadata['source']}")
                text = extract_text_from_eligibility_section_from_link(document.metadata['source'])
                if not text:
                    continue  # Skip if no eligibility text found
                
                questions = extract_questions_from_eligibility_section(text)
                
                if not questions:
                    print("[DEBUG] No questions found using patterns. Generating questions using OpenAI...")
                    questions = generate_questions_from_text(text)
                
                print(f"[DEBUG] Total questions for this document: {len(questions)}")
                print("[DEBUG] Extracted questions from eligibility section:")
                print(questions)
                
                add_to_chroma(questions)
            except Exception as e:
                print(f"[DEBUG] Error processing document: {e}")

    print("[DEBUG] Questions extracted and added to the database successfully.")

def load_documents():
    """Load PDF documents from the specified directory."""
    print(f"[DEBUG] Loading documents from {DATA_PATH}")
    document_loader = PyPDFDirectoryLoader(DATA_PATH)
    documents = document_loader.load()
    print(f"[DEBUG] Loaded {len(documents)} documents")
    return documents

def add_to_chroma(questions):
    """Add extracted questions to the Chroma database."""
    print("[DEBUG] Loading the Chroma database...")
    db = Chroma(persist_directory=CHROMA_PATH, embedding_function=OpenAIEmbeddings(openai_api_key=openai.api_key))
    print("[DEBUG] Chroma database loaded.")

    print("[DEBUG] Adding questions to the database...")
    db.add_texts(questions)
    db.persist()
    print("[DEBUG] Questions added to the database.")

def clear_database():
    """Clear the Chroma database."""
    print(f"[DEBUG] Checking if database path exists at {CHROMA_PATH}")
    if os.path.exists(CHROMA_PATH):
        print("[DEBUG] Database path exists, attempting to remove it...")
        try:
            shutil.rmtree(CHROMA_PATH)
            print(f"[DEBUG] Database cleared at {CHROMA_PATH}")
        except Exception as e:
            print(f"[DEBUG] Error while clearing database: {e}")
    else:
        print("[DEBUG] Database path does not exist, nothing to clear.")

if __name__ == "__main__":
    start_time = datetime.now()
    print(f"[DEBUG] Script started at {start_time}")
    main()
    end_time = datetime.now()
    print(f"[DEBUG] Script ended at {end_time}")
    print(f"[DEBUG] Total execution time: {end_time - start_time}")

